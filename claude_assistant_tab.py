#!/usr/bin/env python3
"""
Enhanced Claude Assistant Tab for Waverider GUI
================================================

Features:
- Natural language interaction with waverider design tool
- Constraint validation before proposing designs
- DOCX report generation for designs
- Off-design and multi-Mach performance analysis
- **NEW: Direct GUI control - can execute actions on the GUI**
"""

import sys
import os
import json
import re
import numpy as np
from datetime import datetime
from pathlib import Path
import tempfile

from PyQt5.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QLabel,
                             QPushButton, QGroupBox, QGridLayout, QFrame,
                             QLineEdit, QTextEdit, QScrollArea, QSplitter,
                             QMessageBox, QInputDialog, QApplication,
                             QSizePolicy, QComboBox, QFileDialog, QProgressBar,
                             QCheckBox)
from PyQt5.QtCore import Qt, pyqtSignal, QThread, QTimer
from PyQt5.QtGui import QFont, QTextCursor, QColor

# Matplotlib for plots
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend for saving
import matplotlib.pyplot as plt
from io import BytesIO

# Try to import anthropic
try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

# Try to import docx generation
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# =============================================================================
# CONSTRAINT VALIDATION
# =============================================================================

class DesignValidator:
    """Validates waverider design parameters against physics constraints."""
    
    @staticmethod
    def get_mach_angle(mach):
        """Get Mach angle in degrees for given Mach number."""
        if mach <= 1:
            return 90.0
        return np.degrees(np.arcsin(1.0 / mach))
    
    @staticmethod
    def get_min_beta(design_mach, margin=1.0):
        """Get minimum valid shock angle for design Mach."""
        mach_angle = DesignValidator.get_mach_angle(design_mach)
        return mach_angle + margin
    
    @staticmethod
    def validate_design(params):
        """
        Validate design parameters against all constraints.
        
        Returns:
            dict: {
                'valid': bool,
                'errors': list of error strings,
                'warnings': list of warning strings,
                'suggestions': list of suggestions
            }
        """
        errors = []
        warnings = []
        suggestions = []
        
        # Extract parameters with defaults
        X1 = params.get('X1', 0.25)
        X2 = params.get('X2', 0.5)
        X3 = params.get('X3', 0.5)
        X4 = params.get('X4', 0.5)
        design_mach = params.get('design_Mach', 5.0)
        design_beta = params.get('design_beta', 20.0)
        width = params.get('width', 2.0)
        height = params.get('height', 1.0)
        flight_mach = params.get('flight_Mach', design_mach)
        
        # Check X1 bounds (0.05 - 0.45)
        if X1 < 0.05:
            errors.append(f"X1 = {X1:.3f} is below minimum (0.05)")
            suggestions.append("Increase X1 to at least 0.05")
        elif X1 > 0.45:
            errors.append(f"X1 = {X1:.3f} is above maximum (0.45)")
            suggestions.append("Decrease X1 to at most 0.45")
        
        # Check X2, X3, X4 bounds (0.1 - 0.9)
        for name, val in [('X2', X2), ('X3', X3), ('X4', X4)]:
            if val < 0.1:
                errors.append(f"{name} = {val:.3f} is below minimum (0.1)")
                suggestions.append(f"Increase {name} to at least 0.1")
            elif val > 0.9:
                errors.append(f"{name} = {val:.3f} is above maximum (0.9)")
                suggestions.append(f"Decrease {name} to at most 0.9")
        
        # Check design Mach (3-6 for training data)
        if design_mach < 3:
            errors.append(f"Design Mach {design_mach} is below training range (3-6)")
            suggestions.append("Use Design Mach 3 or higher")
        elif design_mach > 6:
            errors.append(f"Design Mach {design_mach} is above training range (3-6)")
            suggestions.append("Use Design Mach 6 or lower")
        elif design_mach not in [3, 4, 5, 6]:
            warnings.append(f"Design Mach {design_mach} is not an integer (training used M3, M4, M5, M6)")
            suggestions.append("Consider using integer Mach numbers for best accuracy")
        
        # Check shock angle > Mach angle (CRITICAL)
        min_beta = DesignValidator.get_min_beta(design_mach)
        if design_beta < min_beta:
            errors.append(f"β = {design_beta:.1f}° is below Mach angle limit ({min_beta:.1f}° for M{design_mach:.0f})")
            errors.append("This design is PHYSICALLY IMPOSSIBLE - no attached shock can exist!")
            suggestions.append(f"Increase β to at least {min_beta:.1f}°")
        elif design_beta < min_beta + 2:
            warnings.append(f"β = {design_beta:.1f}° is very close to Mach angle limit ({min_beta:.1f}°)")
            suggestions.append("Consider increasing β for more robust shock attachment")
        
        # Check beta upper bound
        if design_beta > 26:
            warnings.append(f"β = {design_beta:.1f}° is above typical training range (15-26°)")
            suggestions.append("Higher β may result in increased uncertainty")
        
        # Check geometry dimensions
        if width <= 0:
            errors.append(f"Width must be positive (got {width})")
        if height <= 0:
            errors.append(f"Height must be positive (got {height})")
        if width > 0 and height > 0 and height > width:
            warnings.append(f"Height ({height}m) > Width ({width}m) is unusual for waveriders")
        
        # Check off-design flight conditions
        if flight_mach is not None:
            valid_flight_range = {
                3: [3, 4, 5],
                4: [4, 5, 6],
                5: [5, 6, 7],
                6: [6, 7, 8]
            }
            dm_int = int(round(design_mach))
            if dm_int in valid_flight_range:
                if flight_mach < min(valid_flight_range[dm_int]) - 0.5:
                    warnings.append(f"Flight M{flight_mach} below training range for Design M{dm_int}")
                    suggestions.append(f"Training data covers M{min(valid_flight_range[dm_int])}-M{max(valid_flight_range[dm_int])}")
                elif flight_mach > max(valid_flight_range[dm_int]) + 0.5:
                    warnings.append(f"Flight M{flight_mach} above training range for Design M{dm_int}")
                    suggestions.append("High uncertainty expected - validate with CFD")
        
        return {
            'valid': len(errors) == 0,
            'errors': errors,
            'warnings': warnings,
            'suggestions': suggestions
        }
    
    @staticmethod
    def format_validation_report(validation_result):
        """Format validation result as readable text."""
        lines = []
        
        if validation_result['valid']:
            lines.append("✅ DESIGN VALIDATION: PASSED")
        else:
            lines.append("❌ DESIGN VALIDATION: FAILED")
        
        lines.append("")
        
        if validation_result['errors']:
            lines.append("ERRORS:")
            for err in validation_result['errors']:
                lines.append(f"  ❌ {err}")
            lines.append("")
        
        if validation_result['warnings']:
            lines.append("WARNINGS:")
            for warn in validation_result['warnings']:
                lines.append(f"  ⚠️ {warn}")
            lines.append("")
        
        if validation_result['suggestions']:
            lines.append("SUGGESTIONS:")
            for sug in validation_result['suggestions']:
                lines.append(f"  💡 {sug}")
        
        return "\n".join(lines)


# =============================================================================
# GUI ACTION EXECUTOR
# =============================================================================

class GUIActionExecutor:
    """Executes actions on the waverider GUI."""
    
    # Define available actions and their parameters
    AVAILABLE_ACTIONS = {
        'set_design_mach': {
            'description': 'Set the design Mach number',
            'params': {'mach': 'float (1.1-20.0)'},
            'example': '[ACTION: set_design_mach, mach=5.0]'
        },
        'set_shock_angle': {
            'description': 'Set the design shock angle (beta)',
            'params': {'beta': 'float (5.0-89.0) degrees'},
            'example': '[ACTION: set_shock_angle, beta=15.0]'
        },
        'set_geometry': {
            'description': 'Set the waverider physical dimensions',
            'params': {'width': 'float (m)', 'height': 'float (m)'},
            'example': '[ACTION: set_geometry, width=3.0, height=1.34]'
        },
        'set_shape_parameters': {
            'description': 'Set the X1-X4 shape parameters',
            'params': {'X1': 'float (0.05-0.45)', 'X2': 'float (0.1-0.9)', 
                      'X3': 'float (0.1-0.9)', 'X4': 'float (0.1-0.9)'},
            'example': '[ACTION: set_shape_parameters, X1=0.15, X2=0.42, X3=0.50, X4=0.36]'
        },
        'generate_waverider': {
            'description': 'Generate the waverider geometry with current parameters',
            'params': {},
            'example': '[ACTION: generate_waverider]'
        },
        'run_cfd_analysis': {
            'description': 'Run PySAGAS CFD analysis on current geometry',
            'params': {'mach': 'float (optional, default=5.0)', 'aoa': 'float (optional, default=0.0) degrees'},
            'example': '[ACTION: run_cfd_analysis, mach=5.0, aoa=2.0]'
        },
        'set_flight_conditions': {
            'description': 'Set flight conditions in Off-Design NN tab',
            'params': {'flight_mach': 'float', 'flight_aoa': 'float (degrees)'},
            'example': '[ACTION: set_flight_conditions, flight_mach=5.0, flight_aoa=2.0]'
        },
        'run_surrogate_prediction': {
            'description': 'Run neural network surrogate prediction',
            'params': {},
            'example': '[ACTION: run_surrogate_prediction]'
        },
        'run_clcd_hunter': {
            'description': 'Run the CL/CD Hunter optimization',
            'params': {'target_clcd': 'float (optional)'},
            'example': '[ACTION: run_clcd_hunter]'
        },
        'run_multimach_hunter': {
            'description': 'Run Multi-Mach Hunter optimization',
            'params': {'mach_min': 'int', 'mach_max': 'int', 'objective': 'str (robust/mean/consistent/balanced)'},
            'example': '[ACTION: run_multimach_hunter, mach_min=4, mach_max=7, objective=robust]'
        },
        'export_stl': {
            'description': 'Export waverider geometry to STL file',
            'params': {'filename': 'str (optional)'},
            'example': '[ACTION: export_stl, filename=waverider.stl]'
        },
        'switch_tab': {
            'description': 'Switch to a specific tab in the GUI',
            'params': {'tab': 'str (3D View/Aero Analysis/Optimization/Off-Design NN/Multi-Mach)'},
            'example': '[ACTION: switch_tab, tab=Off-Design NN]'
        },
        'get_current_results': {
            'description': 'Get the current analysis results',
            'params': {},
            'example': '[ACTION: get_current_results]'
        }
    }
    
    def __init__(self, parent_gui):
        self.gui = parent_gui
        self.last_result = None
        self.action_log = []
    
    def get_actions_description(self):
        """Get a formatted description of all available actions for the system prompt."""
        lines = ["Available GUI Actions:"]
        for action_name, info in self.AVAILABLE_ACTIONS.items():
            lines.append(f"\n  {action_name}: {info['description']}")
            if info['params']:
                params_str = ", ".join([f"{k}={v}" for k, v in info['params'].items()])
                lines.append(f"    Parameters: {params_str}")
            lines.append(f"    Example: {info['example']}")
        return "\n".join(lines)
    
    def parse_actions(self, response_text):
        """Parse action commands from Claude's response."""
        # Pattern: [ACTION: action_name, param1=value1, param2=value2]
        pattern = r'\[ACTION:\s*(\w+)(?:,\s*([^\]]+))?\]'
        matches = re.findall(pattern, response_text)
        
        actions = []
        for match in matches:
            action_name = match[0]
            params_str = match[1] if len(match) > 1 else ""
            
            # Parse parameters
            params = {}
            if params_str:
                # Split by comma, but handle potential nested values
                param_pairs = re.findall(r'(\w+)\s*=\s*([^,\]]+)', params_str)
                for key, value in param_pairs:
                    # Try to convert to appropriate type
                    value = value.strip()
                    try:
                        if '.' in value:
                            params[key] = float(value)
                        else:
                            params[key] = int(value)
                    except ValueError:
                        params[key] = value  # Keep as string
            
            actions.append({'action': action_name, 'params': params})
        
        return actions
    
    def execute_action(self, action_name, params):
        """Execute a single action on the GUI."""
        result = {'success': False, 'message': '', 'data': None}
        
        try:
            if action_name == 'set_design_mach':
                mach = params.get('mach', 5.0)
                if hasattr(self.gui, 'm_inf_spin'):
                    self.gui.m_inf_spin.setValue(float(mach))
                    result = {'success': True, 'message': f'Set design Mach to {mach}'}
                    
            elif action_name == 'set_shock_angle':
                beta = params.get('beta', 15.0)
                if hasattr(self.gui, 'beta_spin'):
                    self.gui.beta_spin.setValue(float(beta))
                    result = {'success': True, 'message': f'Set shock angle to {beta}°'}
                    
            elif action_name == 'set_geometry':
                width = params.get('width')
                height = params.get('height')
                if width and hasattr(self.gui, 'width_spin'):
                    self.gui.width_spin.setValue(float(width))
                if height and hasattr(self.gui, 'height_spin'):
                    self.gui.height_spin.setValue(float(height))
                result = {'success': True, 'message': f'Set geometry: width={width}m, height={height}m'}
                
            elif action_name == 'set_shape_parameters':
                msgs = []
                for param in ['X1', 'X2', 'X3', 'X4']:
                    if param in params:
                        slider_name = f'{param.lower()}_slider'
                        spin_name = f'{param.lower()}_spin'
                        value = float(params[param])
                        
                        if hasattr(self.gui, slider_name):
                            slider = getattr(self.gui, slider_name)
                            slider.setValue(int(value * 1000))
                            msgs.append(f'{param}={value:.3f}')
                        elif hasattr(self.gui, spin_name):
                            spin = getattr(self.gui, spin_name)
                            spin.setValue(value)
                            msgs.append(f'{param}={value:.3f}')
                
                result = {'success': True, 'message': f'Set shape parameters: {", ".join(msgs)}'}
                
            elif action_name == 'generate_waverider':
                if hasattr(self.gui, 'generate_waverider'):
                    self.gui.generate_waverider()
                    result = {'success': True, 'message': 'Waverider geometry generated'}
                elif hasattr(self.gui, 'generate_btn'):
                    self.gui.generate_btn.click()
                    result = {'success': True, 'message': 'Waverider geometry generated'}
                    
            elif action_name == 'run_cfd_analysis':
                mach = params.get('mach', 5.0)
                aoa = params.get('aoa', 0.0)
                if hasattr(self.gui, 'analysis_mach_spin'):
                    self.gui.analysis_mach_spin.setValue(float(mach))
                if hasattr(self.gui, 'aoa_spin'):
                    self.gui.aoa_spin.setValue(float(aoa))
                if hasattr(self.gui, 'run_analysis'):
                    self.gui.run_analysis()
                    result = {'success': True, 'message': f'CFD analysis started at M={mach}, α={aoa}°'}
                elif hasattr(self.gui, 'run_analysis_btn'):
                    self.gui.run_analysis_btn.click()
                    result = {'success': True, 'message': f'CFD analysis started at M={mach}, α={aoa}°'}
                    
            elif action_name == 'set_flight_conditions':
                flight_mach = params.get('flight_mach', 5.0)
                flight_aoa = params.get('flight_aoa', 2.0)
                
                if hasattr(self.gui, 'offdesign_tab'):
                    od_tab = self.gui.offdesign_tab
                    if hasattr(od_tab, 'flight_mach_spin'):
                        od_tab.flight_mach_spin.setValue(float(flight_mach))
                    if hasattr(od_tab, 'flight_aoa_spin'):
                        od_tab.flight_aoa_spin.setValue(float(flight_aoa))
                    result = {'success': True, 'message': f'Set flight conditions: M={flight_mach}, α={flight_aoa}°'}
                else:
                    result = {'success': False, 'message': 'Off-Design NN tab not available'}
                    
            elif action_name == 'run_surrogate_prediction':
                if hasattr(self.gui, 'offdesign_tab'):
                    od_tab = self.gui.offdesign_tab
                    if hasattr(od_tab, 'predict_btn'):
                        od_tab.predict_btn.click()
                        result = {'success': True, 'message': 'Surrogate prediction executed'}
                    elif hasattr(od_tab, 'run_prediction'):
                        od_tab.run_prediction()
                        result = {'success': True, 'message': 'Surrogate prediction executed'}
                else:
                    result = {'success': False, 'message': 'Off-Design NN tab not available'}
                    
            elif action_name == 'run_clcd_hunter':
                if hasattr(self.gui, 'offdesign_tab'):
                    od_tab = self.gui.offdesign_tab
                    if hasattr(od_tab, 'run_hunter_btn'):
                        od_tab.run_hunter_btn.click()
                        result = {'success': True, 'message': 'CL/CD Hunter started'}
                    elif hasattr(od_tab, 'run_clcd_hunter'):
                        od_tab.run_clcd_hunter()
                        result = {'success': True, 'message': 'CL/CD Hunter started'}
                else:
                    result = {'success': False, 'message': 'Off-Design NN tab not available'}
                    
            elif action_name == 'run_multimach_hunter':
                if hasattr(self.gui, 'multimach_tab'):
                    mm_tab = self.gui.multimach_tab
                    
                    # Set parameters if provided
                    if 'mach_min' in params and hasattr(mm_tab, 'mach_min_spin'):
                        mm_tab.mach_min_spin.setValue(int(params['mach_min']))
                    if 'mach_max' in params and hasattr(mm_tab, 'mach_max_spin'):
                        mm_tab.mach_max_spin.setValue(int(params['mach_max']))
                    if 'objective' in params and hasattr(mm_tab, 'objective_combo'):
                        obj_map = {'robust': 0, 'mean': 1, 'consistent': 2, 'balanced': 3}
                        idx = obj_map.get(params['objective'].lower(), 0)
                        mm_tab.objective_combo.setCurrentIndex(idx)
                    
                    if hasattr(mm_tab, 'run_hunter_btn'):
                        mm_tab.run_hunter_btn.click()
                        result = {'success': True, 'message': 'Multi-Mach Hunter started'}
                else:
                    result = {'success': False, 'message': 'Multi-Mach tab not available'}
                    
            elif action_name == 'export_stl':
                filename = params.get('filename', 'waverider.stl')
                if hasattr(self.gui, 'export_cad'):
                    self.gui.export_cad()
                    result = {'success': True, 'message': f'Export dialog opened'}
                elif hasattr(self.gui, 'export_btn'):
                    self.gui.export_btn.click()
                    result = {'success': True, 'message': f'Export dialog opened'}
                    
            elif action_name == 'switch_tab':
                tab_name = params.get('tab', '3D View')
                if hasattr(self.gui, 'tab_widget'):
                    tab_map = {
                        '3d view': 0, '3d': 0,
                        'base plane': 1,
                        'leading edge': 2,
                        'geometry schematic': 3, 'schematic': 3,
                        'aero analysis': 4, 'aero': 4, 'analysis': 4,
                        'optimization': 5, 'opt': 5,
                        'surrogate opt': 6, 'surrogate': 6,
                        'off-design nn': 7, 'off-design': 7, 'offdesign': 7,
                        'multi-mach': 8, 'multimach': 8,
                    }
                    idx = tab_map.get(tab_name.lower(), 0)
                    self.gui.tab_widget.setCurrentIndex(idx)
                    result = {'success': True, 'message': f'Switched to tab: {tab_name}'}
                    
            elif action_name == 'get_current_results':
                data = {}
                
                # Get design parameters
                if hasattr(self.gui, 'm_inf_spin'):
                    data['design_Mach'] = self.gui.m_inf_spin.value()
                if hasattr(self.gui, 'beta_spin'):
                    data['design_beta'] = self.gui.beta_spin.value()
                if hasattr(self.gui, 'width_spin'):
                    data['width'] = self.gui.width_spin.value()
                if hasattr(self.gui, 'height_spin'):
                    data['height'] = self.gui.height_spin.value()
                
                # Get analysis parameters
                if hasattr(self.gui, 'analysis_mach_spin'):
                    data['analysis_Mach'] = self.gui.analysis_mach_spin.value()
                if hasattr(self.gui, 'aoa_spin'):
                    data['analysis_AoA'] = self.gui.aoa_spin.value()
                
                # Get shape parameters
                for i, param in enumerate(['X1', 'X2', 'X3', 'X4'], 1):
                    slider_name = f'x{i}_slider'
                    if hasattr(self.gui, slider_name):
                        data[param] = getattr(self.gui, slider_name).value() / 1000.0
                
                # Get last analysis results if available
                if hasattr(self.gui, 'last_analysis_results'):
                    data['analysis'] = self.gui.last_analysis_results
                
                result = {'success': True, 'message': 'Retrieved current state', 'data': data}
                
            else:
                result = {'success': False, 'message': f'Unknown action: {action_name}'}
                
        except Exception as e:
            result = {'success': False, 'message': f'Error executing {action_name}: {str(e)}'}
        
        # Log the action
        self.action_log.append({
            'action': action_name,
            'params': params,
            'result': result,
            'timestamp': datetime.now().isoformat()
        })
        self.last_result = result
        
        return result
    
    def execute_actions(self, actions):
        """Execute a list of actions and return results."""
        results = []
        for action_info in actions:
            result = self.execute_action(action_info['action'], action_info['params'])
            results.append(result)
            
            # Small delay between actions for GUI to update
            QApplication.processEvents()
        
        return results


# =============================================================================
# REPORT GENERATOR
# =============================================================================

class ReportGenerator:
    """Generates DOCX reports for waverider designs."""
    
    def __init__(self, surrogate_models=None):
        self.models = surrogate_models or {}
        
    def generate_offdesign_plot(self, params, mach_range, aoa=2.0):
        """Generate off-design performance plot."""
        if 'CL_CD' not in self.models:
            return None
        
        fig, ax = plt.subplots(figsize=(8, 5))
        
        machs = np.linspace(mach_range[0], mach_range[1], 20)
        clcd_values = []
        clcd_stds = []
        
        for m in machs:
            X = np.array([[
                params['X1'], params['X2'], params['X3'], params['X4'],
                params['design_Mach'], params['design_beta'],
                m, aoa,
                params['width'], params['height']
            ]])
            pred, std = self.models['CL_CD'].predict(X, return_std=True)
            clcd_values.append(pred[0])
            clcd_stds.append(std[0])
        
        clcd_values = np.array(clcd_values)
        clcd_stds = np.array(clcd_stds)
        
        ax.fill_between(machs, clcd_values - clcd_stds, clcd_values + clcd_stds,
                       alpha=0.3, color='blue', label='Uncertainty (±1σ)')
        ax.plot(machs, clcd_values, 'b-', linewidth=2, label='CL/CD Prediction')
        ax.axvline(params['design_Mach'], color='red', linestyle='--', 
                  label=f'Design Mach ({params["design_Mach"]:.0f})')
        
        ax.set_xlabel('Flight Mach Number', fontsize=12)
        ax.set_ylabel('CL/CD', fontsize=12)
        ax.set_title('Off-Design Performance Envelope', fontsize=14)
        ax.legend(loc='best')
        ax.grid(True, alpha=0.3)
        
        buf = BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format='png', dpi=150)
        plt.close(fig)
        buf.seek(0)
        return buf
    
    def generate_report(self, params, output_path, include_offdesign=True, 
                       include_multimach=True, mach_range=(3, 8), aoa=2.0):
        """Generate a comprehensive DOCX report for a waverider design."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = DocxDocument()
        
        # Title
        title = doc.add_heading('Waverider Design Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle with date
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        run.italic = True
        
        doc.add_paragraph()
        
        # Validate design
        validation = DesignValidator.validate_design(params)
        
        # Executive Summary
        doc.add_heading('1. Executive Summary', level=1)
        if validation['valid']:
            doc.add_paragraph(
                f"This report presents a waverider design optimized for Mach {params.get('design_Mach', 5):.0f} "
                f"with a shock angle of {params.get('design_beta', 20):.1f}°."
            )
        else:
            p = doc.add_paragraph()
            p.add_run("⚠️ WARNING: ").bold = True
            p.add_run("This design has validation errors.")
        
        # Design Parameters
        doc.add_heading('2. Design Parameters', level=1)
        
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        
        params_data = [
            ('Design Mach', f"{params.get('design_Mach', 5):.1f}"),
            ('Shock Angle (β)', f"{params.get('design_beta', 20):.2f}°"),
            ('X1', f"{params.get('X1', 0.25):.4f}"),
            ('X2', f"{params.get('X2', 0.5):.4f}"),
            ('X3', f"{params.get('X3', 0.5):.4f}"),
            ('X4', f"{params.get('X4', 0.5):.4f}"),
            ('Width', f"{params.get('width', 2):.3f} m"),
            ('Height', f"{params.get('height', 1):.3f} m"),
        ]
        for i, (label, value) in enumerate(params_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
        
        # Validation Results
        doc.add_heading('3. Validation Results', level=1)
        if validation['valid']:
            doc.add_paragraph("✅ All validation checks passed.")
        else:
            doc.add_paragraph("❌ Validation failed:")
            for err in validation['errors']:
                doc.add_paragraph(f"• {err}", style='List Bullet')
        
        # Save document
        doc.save(output_path)
        return output_path


# =============================================================================
# SYSTEM PROMPT WITH ACTIONS
# =============================================================================

def get_system_prompt(action_executor):
    """Generate system prompt with available actions."""
    
    actions_desc = action_executor.get_actions_description() if action_executor else ""
    
    return f"""You are an expert aerospace engineer assistant integrated into a waverider design GUI. You can both provide advice AND directly control the GUI to execute tasks.

**CRITICAL: You can execute actions on the GUI!**

When you want to perform an action, include an action command in your response using this format:
[ACTION: action_name, param1=value1, param2=value2]

The action will be executed automatically after your response.

{actions_desc}

**IMPORTANT RULES FOR ACTIONS:**
1. Always validate designs before setting parameters (check β > Mach angle)
2. You can chain multiple actions in one response
3. After setting parameters, call [ACTION: generate_waverider] to see the result
4. Use [ACTION: run_surrogate_prediction] to get performance estimates
5. Explain what you're doing as you execute actions

**Physics Constraints (MUST satisfy):**
- Shock angle β must EXCEED the Mach angle: β > arcsin(1/M) + 1°
  - Mach 3: β > 20.5°
  - Mach 4: β > 15.5°
  - Mach 5: β > 12.5°
  - Mach 6: β > 10.6°

**Parameter Bounds:**
- X1: 0.05 - 0.45 (flat region length)
- X2: 0.1 - 0.9 (shockwave height - MOST INFLUENTIAL)
- X3: 0.1 - 0.9 (central upper surface)
- X4: 0.1 - 0.9 (side upper surface)

**Example Interaction:**
User: "Design a Mach 5 waverider with good CL/CD"
Assistant: "I'll create a Mach 5 design with optimized parameters for good lift-to-drag ratio.

Setting up the design:
[ACTION: set_design_mach, mach=5.0]
[ACTION: set_shock_angle, beta=15.0]
[ACTION: set_shape_parameters, X1=0.15, X2=0.42, X3=0.50, X4=0.36]
[ACTION: set_geometry, width=2.7, height=0.64]
[ACTION: generate_waverider]

Now let me run the surrogate prediction to check performance:
[ACTION: set_flight_conditions, flight_mach=5.0, flight_aoa=2.0]
[ACTION: run_surrogate_prediction]

This design should give you a CL/CD around 5.0-5.5 at the design point."
"""


# =============================================================================
# CLAUDE WORKER
# =============================================================================

class ClaudeWorker(QThread):
    """Worker thread for Claude API calls."""
    
    response_ready = pyqtSignal(str)
    error_occurred = pyqtSignal(str)
    stream_chunk = pyqtSignal(str)
    
    def __init__(self, api_key, messages, system_prompt, context=""):
        super().__init__()
        self.api_key = api_key
        self.messages = messages
        self.system_prompt = system_prompt
        self.context = context
        
    def run(self):
        try:
            client = anthropic.Anthropic(api_key=self.api_key)
            
            full_system = self.system_prompt
            if self.context:
                full_system += f"\n\n**Current GUI State:**\n{self.context}"
            
            with client.messages.stream(
                model="claude-sonnet-4-20250514",
                max_tokens=2000,
                system=full_system,
                messages=self.messages
            ) as stream:
                full_response = ""
                for text in stream.text_stream:
                    full_response += text
                    self.stream_chunk.emit(text)
                
                self.response_ready.emit(full_response)
                
        except anthropic.AuthenticationError:
            self.error_occurred.emit("Invalid API key. Please check your Anthropic API key.")
        except anthropic.RateLimitError:
            self.error_occurred.emit("Rate limit exceeded. Please wait a moment and try again.")
        except Exception as e:
            self.error_occurred.emit(f"API Error: {str(e)}")


# =============================================================================
# CHAT MESSAGE WIDGET
# =============================================================================

class ChatMessage(QFrame):
    """Widget for displaying a single chat message."""
    
    def __init__(self, role, content, parent=None):
        super().__init__(parent)
        self.role = role
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 5, 10, 5)
        
        role_label = QLabel("You:" if role == "user" else "Claude:")
        role_label.setFont(QFont("Arial", 9, QFont.Bold))
        role_label.setStyleSheet(f"color: {'#F59E0B' if role == 'user' else '#4ADE80'};")
        layout.addWidget(role_label)
        
        self.content_label = QTextEdit()
        self.content_label.setReadOnly(True)
        self.content_label.setPlainText(content)
        self.content_label.setFont(QFont("Arial", 10))
        self.content_label.setFrameStyle(QFrame.NoFrame)
        self.content_label.setStyleSheet("QTextEdit { background-color: transparent; border: none; color: #FFFFFF; }")
        self.content_label.document().setTextWidth(self.content_label.viewport().width())
        height = int(self.content_label.document().size().height()) + 10
        self.content_label.setMinimumHeight(min(height, 300))
        self.content_label.setMaximumHeight(500)
        
        layout.addWidget(self.content_label)
        
        bg_color = "#78350F" if role == "user" else "#1A1A1A"
        self.setStyleSheet(f"ChatMessage {{ background-color: {bg_color}; border-radius: 10px; margin: 5px; }}")
    
    def append_text(self, text):
        current = self.content_label.toPlainText()
        self.content_label.setPlainText(current + text)
        cursor = self.content_label.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.content_label.setTextCursor(cursor)


class ActionResultWidget(QFrame):
    """Widget for displaying action execution results."""
    
    def __init__(self, results, parent=None):
        super().__init__(parent)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 5, 10, 5)
        
        header = QLabel("🔧 Actions Executed:")
        header.setFont(QFont("Arial", 9, QFont.Bold))
        header.setStyleSheet("color: #888888;")
        layout.addWidget(header)
        
        for result in results:
            icon = "✅" if result['success'] else "❌"
            msg = QLabel(f"  {icon} {result['message']}")
            msg.setFont(QFont("Arial", 9))
            msg.setStyleSheet("color: #FFFFFF;")
            layout.addWidget(msg)
        
        self.setStyleSheet("ActionResultWidget { background-color: #1A1A1A; border-radius: 8px; margin: 3px; border: 1px solid #333333; }")


# =============================================================================
# MAIN TAB WIDGET
# =============================================================================

class ClaudeAssistantTab(QWidget):
    """Enhanced Claude Assistant Tab with GUI control capabilities."""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent_gui = parent
        
        self.api_key = os.environ.get('ANTHROPIC_API_KEY', '')
        self.messages = []
        self.message_widgets = []
        
        # Surrogate models
        self.ensemble_CL = None
        self.ensemble_CD = None
        self.ensemble_CL_CD = None
        self.ensemble_Volume = None
        self.models_loaded = False
        
        # Tools
        self.validator = DesignValidator()
        self.report_generator = None
        self.action_executor = GUIActionExecutor(parent) if parent else None
        
        self.worker = None
        self.current_assistant_message = None
        
        # Auto-execute actions flag
        self.auto_execute_actions = True
        
        self.init_ui()
        
    def init_ui(self):
        """Initialize the user interface."""
        main_layout = QVBoxLayout(self)
        
        # API Key Section
        api_group = QGroupBox("🔑 API Configuration")
        api_layout = QHBoxLayout(api_group)
        
        api_layout.addWidget(QLabel("API Key:"))
        self.api_key_input = QLineEdit()
        self.api_key_input.setEchoMode(QLineEdit.Password)
        self.api_key_input.setPlaceholderText("Enter your Anthropic API key (sk-ant-...)")
        if self.api_key:
            self.api_key_input.setText(self.api_key)
        api_layout.addWidget(self.api_key_input)
        
        save_key_btn = QPushButton("Save")
        save_key_btn.clicked.connect(self.save_api_key)
        api_layout.addWidget(save_key_btn)
        
        self.api_status = QLabel("❌ Not configured")
        self.api_status.setStyleSheet("color: #EF4444;")
        api_layout.addWidget(self.api_status)
        
        main_layout.addWidget(api_group)
        
        if self.api_key:
            self.api_status.setText("✅ Key configured")
            self.api_status.setStyleSheet("color: #4ADE80;")
        
        # Main splitter
        splitter = QSplitter(Qt.Horizontal)
        
        # Left panel - Quick actions
        left_panel = self.create_quick_actions_panel()
        splitter.addWidget(left_panel)
        
        # Right panel - Chat
        chat_panel = self.create_chat_panel()
        splitter.addWidget(chat_panel)
        
        splitter.setSizes([280, 520])
        main_layout.addWidget(splitter)
        
        if not ANTHROPIC_AVAILABLE:
            self.show_library_warning()
    
    def create_quick_actions_panel(self):
        """Create panel with quick action buttons."""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        # Settings
        settings_group = QGroupBox("⚙️ Settings")
        settings_layout = QVBoxLayout(settings_group)
        
        self.auto_execute_check = QCheckBox("Auto-execute actions")
        self.auto_execute_check.setChecked(True)
        self.auto_execute_check.setToolTip("Automatically execute GUI actions from Claude's responses")
        self.auto_execute_check.toggled.connect(lambda x: setattr(self, 'auto_execute_actions', x))
        settings_layout.addWidget(self.auto_execute_check)
        
        layout.addWidget(settings_group)
        
        # Quick Actions
        actions_group = QGroupBox("⚡ Quick Actions")
        actions_layout = QVBoxLayout(actions_group)
        
        quick_actions = [
            ("🎯 Design Waverider", "Design an optimal Mach 5 waverider with good CL/CD. Set the parameters, generate the geometry, and run the prediction."),
            ("📊 Run Prediction", "Run the surrogate prediction for the current design parameters and tell me the results."),
            ("🔬 Analyze & Optimize", "Analyze my current design and suggest improvements. Then apply the improvements and show me the new performance."),
            ("🌐 Multi-Mach Analysis", "Run a multi-Mach analysis for my current design across Mach 4-7 and show me the performance trade-offs."),
            ("✅ Validate Design", "Validate my current design parameters against all physics constraints."),
        ]
        
        for label, prompt in quick_actions:
            btn = QPushButton(label)
            btn.clicked.connect(lambda checked, p=prompt: self.send_quick_action(p))
            btn.setStyleSheet("""
                QPushButton {
                    text-align: left;
                    padding: 8px;
                    border: 1px solid #333333;
                    border-radius: 5px;
                    background-color: #1A1A1A;
                    color: #FFFFFF;
                }
                QPushButton:hover {
                    background-color: #78350F;
                    border-color: #F59E0B;
                }
            """)
            actions_layout.addWidget(btn)
        
        layout.addWidget(actions_group)
        
        # Report Generation
        report_group = QGroupBox("📄 Report Generation")
        report_layout = QVBoxLayout(report_group)
        
        gen_report_btn = QPushButton("📝 Generate Design Report")
        gen_report_btn.clicked.connect(self.generate_report)
        gen_report_btn.setStyleSheet("""
            QPushButton {
                padding: 10px;
                background-color: #78350F;
                color: #F59E0B;
                border-radius: 5px;
                font-weight: bold;
                border: 1px solid #F59E0B;
            }
            QPushButton:hover { background-color: #F59E0B; color: #0A0A0A; }
        """)
        report_layout.addWidget(gen_report_btn)
        
        layout.addWidget(report_group)
        
        # Context info
        context_group = QGroupBox("📋 Current Context")
        context_layout = QVBoxLayout(context_group)
        
        self.context_display = QTextEdit()
        self.context_display.setReadOnly(True)
        self.context_display.setMaximumHeight(150)
        self.context_display.setFont(QFont("Courier", 8))
        context_layout.addWidget(self.context_display)
        
        refresh_btn = QPushButton("🔄 Refresh")
        refresh_btn.clicked.connect(self.refresh_context)
        context_layout.addWidget(refresh_btn)
        
        layout.addWidget(context_group)
        
        layout.addStretch()
        
        return panel
    
    def create_chat_panel(self):
        """Create the main chat panel."""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        self.chat_scroll = QScrollArea()
        self.chat_scroll.setWidgetResizable(True)
        self.chat_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        
        self.chat_container = QWidget()
        self.chat_layout = QVBoxLayout(self.chat_container)
        self.chat_layout.addStretch()
        
        self.chat_scroll.setWidget(self.chat_container)
        layout.addWidget(self.chat_scroll)
        
        # Welcome message
        welcome = QLabel(
            "👋 Welcome! I'm your waverider design assistant.\n\n"
            "🚀 NEW: I can now CONTROL the GUI directly!\n"
            "• Ask me to design a waverider and I'll set the parameters\n"
            "• Ask me to run predictions and I'll execute them\n"
            "• Ask me to optimize and I'll apply changes automatically\n\n"
            "Try: \"Design a Mach 5 waverider with good CL/CD\""
        )
        welcome.setWordWrap(True)
        welcome.setStyleSheet("""
            QLabel {
                background-color: #1A1A1A;
                padding: 15px;
                border-radius: 10px;
                color: #4ADE80;
            }
        """)
        self.chat_layout.insertWidget(0, welcome)
        
        # Input area
        input_layout = QHBoxLayout()
        
        self.chat_input = QLineEdit()
        self.chat_input.setPlaceholderText("Ask me to design, analyze, or optimize...")
        self.chat_input.setFont(QFont("Arial", 11))
        self.chat_input.returnPressed.connect(self.send_message)
        self.chat_input.setStyleSheet("""
            QLineEdit {
                padding: 10px;
                border: 2px solid #333333;
                border-radius: 20px;
                background-color: #1A1A1A;
                color: #FFFFFF;
            }
            QLineEdit:focus { border-color: #F59E0B; }
        """)
        input_layout.addWidget(self.chat_input)
        
        self.send_btn = QPushButton("Send")
        self.send_btn.clicked.connect(self.send_message)
        self.send_btn.setStyleSheet("""
            QPushButton {
                background-color: #F59E0B;
                color: #0A0A0A;
                padding: 10px 20px;
                border-radius: 20px;
                font-weight: bold;
            }
            QPushButton:hover { background-color: #78350F; color: #F59E0B; }
            QPushButton:disabled { background-color: #333333; color: #888888; }
        """)
        input_layout.addWidget(self.send_btn)
        
        layout.addLayout(input_layout)
        
        # Clear button
        clear_layout = QHBoxLayout()
        clear_layout.addStretch()
        clear_btn = QPushButton("🗑️ Clear Chat")
        clear_btn.clicked.connect(self.clear_chat)
        clear_btn.setStyleSheet("QPushButton { color: #EF4444; border: none; } QPushButton:hover { text-decoration: underline; }")
        clear_layout.addWidget(clear_btn)
        layout.addLayout(clear_layout)
        
        return panel
    
    def show_library_warning(self):
        warning = QMessageBox(self)
        warning.setIcon(QMessageBox.Warning)
        warning.setWindowTitle("Library Not Found")
        warning.setText("The 'anthropic' Python library is not installed.")
        warning.setInformativeText("Install with: pip install anthropic")
        warning.exec_()
    
    def save_api_key(self):
        key = self.api_key_input.text().strip()
        if key:
            self.api_key = key
            self.api_status.setText("✅ Key configured")
            self.api_status.setStyleSheet("color: #4ADE80;")
            QMessageBox.information(self, "Saved", "API key saved for this session.")
        else:
            self.api_status.setText("❌ Not configured")
            self.api_status.setStyleSheet("color: #EF4444;")
    
    def get_gui_context(self):
        """Get current state from the main GUI."""
        context = {}
        
        if self.parent_gui:
            try:
                # Design parameters
                if hasattr(self.parent_gui, 'm_inf_spin'):
                    context['design_Mach'] = self.parent_gui.m_inf_spin.value()
                if hasattr(self.parent_gui, 'beta_spin'):
                    context['design_beta'] = self.parent_gui.beta_spin.value()
                if hasattr(self.parent_gui, 'width_spin'):
                    context['width'] = self.parent_gui.width_spin.value()
                if hasattr(self.parent_gui, 'height_spin'):
                    context['height'] = self.parent_gui.height_spin.value()
                
                # Shape parameters
                for i in range(1, 5):
                    slider_name = f'x{i}_slider'
                    if hasattr(self.parent_gui, slider_name):
                        context[f'X{i}'] = getattr(self.parent_gui, slider_name).value() / 1000.0
                
                # Flight conditions from Off-Design tab
                if hasattr(self.parent_gui, 'offdesign_tab'):
                    od_tab = self.parent_gui.offdesign_tab
                    if hasattr(od_tab, 'flight_mach_spin'):
                        context['flight_Mach'] = od_tab.flight_mach_spin.value()
                    if hasattr(od_tab, 'flight_aoa_spin'):
                        context['flight_AoA'] = od_tab.flight_aoa_spin.value()
                    
                    if od_tab.models_loaded:
                        context['surrogate_loaded'] = True
                        self.models_loaded = True
                        self.ensemble_CL_CD = od_tab.ensemble_CL_CD
                        
            except Exception as e:
                context['error'] = str(e)
        
        return context
    
    def refresh_context(self):
        """Refresh and display the current context."""
        context = self.get_gui_context()
        validation = DesignValidator.validate_design(context)
        
        text = f"Design: M{context.get('design_Mach', '?')}, β={context.get('design_beta', '?')}°\n"
        text += f"Size: {context.get('width', '?')}×{context.get('height', '?')} m\n"
        text += f"X1-X4: [{context.get('X1', 0):.2f}, {context.get('X2', 0):.2f}, "
        text += f"{context.get('X3', 0):.2f}, {context.get('X4', 0):.2f}]\n"
        text += f"Flight: M{context.get('flight_Mach', '?')}, α={context.get('flight_AoA', '?')}°\n"
        text += f"Valid: {'✅' if validation['valid'] else '❌'}"
        text += f" | Surrogate: {'✅' if context.get('surrogate_loaded') else '❌'}"
        
        self.context_display.setText(text)
        return context
    
    def build_context_string(self):
        """Build context string for Claude."""
        context = self.get_gui_context()
        validation = DesignValidator.validate_design(context)
        
        ctx_str = f"""Design Mach: {context.get('design_Mach', 'unknown')}
Design Beta: {context.get('design_beta', 'unknown')}°
Width: {context.get('width', 'unknown')} m
Height: {context.get('height', 'unknown')} m
X1: {context.get('X1', 'unknown')}
X2: {context.get('X2', 'unknown')}
X3: {context.get('X3', 'unknown')}
X4: {context.get('X4', 'unknown')}
Flight Mach: {context.get('flight_Mach', 'unknown')}
Flight AoA: {context.get('flight_AoA', 'unknown')}°
Validation: {'VALID' if validation['valid'] else 'INVALID - ' + '; '.join(validation['errors'][:2])}
Surrogate Loaded: {context.get('surrogate_loaded', False)}"""
        
        return ctx_str
    
    def send_quick_action(self, prompt):
        self.chat_input.setText(prompt)
        self.send_message()
    
    def send_message(self):
        if not ANTHROPIC_AVAILABLE:
            self.show_library_warning()
            return
        
        if not self.api_key:
            QMessageBox.warning(self, "API Key Required", "Please enter your Anthropic API key first.")
            return
        
        message = self.chat_input.text().strip()
        if not message:
            return
        
        self.chat_input.clear()
        
        # Add user message
        user_widget = ChatMessage("user", message)
        self.chat_layout.insertWidget(self.chat_layout.count() - 1, user_widget)
        self.message_widgets.append(user_widget)
        
        self.messages.append({"role": "user", "content": message})
        
        # Placeholder for response
        self.current_assistant_message = ChatMessage("assistant", "")
        self.chat_layout.insertWidget(self.chat_layout.count() - 1, self.current_assistant_message)
        self.message_widgets.append(self.current_assistant_message)
        
        QTimer.singleShot(100, self.scroll_to_bottom)
        
        self.chat_input.setEnabled(False)
        self.send_btn.setEnabled(False)
        
        context = self.build_context_string()
        system_prompt = get_system_prompt(self.action_executor)
        
        self.worker = ClaudeWorker(self.api_key, self.messages.copy(), system_prompt, context)
        self.worker.stream_chunk.connect(self.on_stream_chunk)
        self.worker.response_ready.connect(self.on_response_ready)
        self.worker.error_occurred.connect(self.on_error)
        self.worker.start()
    
    def on_stream_chunk(self, chunk):
        if self.current_assistant_message:
            self.current_assistant_message.append_text(chunk)
            self.scroll_to_bottom()
    
    def on_response_ready(self, response):
        self.messages.append({"role": "assistant", "content": response})
        
        # Parse and execute actions if enabled
        if self.auto_execute_actions and self.action_executor:
            actions = self.action_executor.parse_actions(response)
            if actions:
                # Execute actions
                results = self.action_executor.execute_actions(actions)
                
                # Show action results
                result_widget = ActionResultWidget(results)
                self.chat_layout.insertWidget(self.chat_layout.count() - 1, result_widget)
                self.message_widgets.append(result_widget)
                
                # Refresh context display
                self.refresh_context()
        
        self.chat_input.setEnabled(True)
        self.send_btn.setEnabled(True)
        self.chat_input.setFocus()
        self.scroll_to_bottom()
    
    def on_error(self, error_message):
        if self.current_assistant_message:
            self.current_assistant_message.content_label.setPlainText(f"❌ Error: {error_message}")
        self.chat_input.setEnabled(True)
        self.send_btn.setEnabled(True)
    
    def scroll_to_bottom(self):
        scrollbar = self.chat_scroll.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())
    
    def clear_chat(self):
        for widget in self.message_widgets:
            widget.deleteLater()
        self.message_widgets.clear()
        self.messages.clear()
        
        welcome = QLabel("👋 Chat cleared! What would you like me to do?")
        welcome.setWordWrap(True)
        welcome.setStyleSheet("QLabel { background-color: #1A1A1A; padding: 15px; border-radius: 10px; color: #4ADE80; }")
        self.chat_layout.insertWidget(0, welcome)
    
    def generate_report(self):
        """Generate a DOCX report for the current design."""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Missing Library",
                              "python-docx not installed.\n\nRun: pip install python-docx")
            return
        
        context = self.get_gui_context()
        
        if not context:
            QMessageBox.warning(self, "No Design", "Please set up a design first.")
            return
        
        if self.report_generator is None:
            self.report_generator = ReportGenerator()
        
        default_name = f"waverider_report_M{context.get('design_Mach', 5):.0f}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath, _ = QFileDialog.getSaveFileName(
            self, "Save Report", default_name, "Word Documents (*.docx)"
        )
        
        if not filepath:
            return
        
        try:
            self.report_generator.generate_report(context, filepath)
            QMessageBox.information(self, "Report Generated", f"Report saved to:\n{filepath}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate report:\n{str(e)}")


# For standalone testing
if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = ClaudeAssistantTab()
    window.setWindowTitle("Enhanced Claude Assistant")
    window.resize(950, 750)
    window.show()
    sys.exit(app.exec_())
